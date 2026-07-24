"""Unit coverage for :mod:`ragsage.parsing.docx` — the DOCX parse path.

DOCX carries real structure, so the contract is to *honour* it: heading styles
become Markdown headings, list nesting becomes indented list items, and table
grids become GitHub pipe tables, all in document order. The fixtures are built
programmatically with ``python-docx`` (already a dependency) into an in-memory
``.docx``, so the tests exercise the same library the parser reads back.

Two layers are checked: the Markdown :func:`parse_docx` emits directly (single
layout-free page, heading levels, indented lists, pipe tables), and the chunk-level
behaviour once the real backend runs — heading context in ``metadata["headings"]``
and whole tables surviving into a single chunk.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ragsage.models import Page, RawSource
from ragsage.parsing import HeuristicBackend
from ragsage.parsing.docx import parse_docx

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# --------------------------------------------------------------------------- #
# Fixture builders — a real .docx, in memory
# --------------------------------------------------------------------------- #


def _set_numbering(paragraph: object, level: int) -> None:
    """Attach a ``<w:numPr>`` with the given 0-based ``<w:ilvl>`` to a paragraph.

    python-docx's built-in list *styles* don't inject numbering (that lives in the
    template's numbering.xml), so to exercise the authoritative numPr/ilvl path we
    write the element directly — the same shape Word produces for a nested list.
    """
    p_pr = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def _docx_bytes(build: object) -> bytes:
    """Build a document with ``build(document)`` and return its ``.docx`` bytes."""
    document = DocxDocument()
    build(document)  # type: ignore[operator]
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _structured_docx() -> bytes:
    """Headings + a nested bulleted list + a table, interleaved in document order."""

    def build(document: object) -> None:
        document.add_heading("Quarterly Report", level=1)  # type: ignore[attr-defined]
        document.add_paragraph("Opening remarks for the quarter.")  # type: ignore[attr-defined]
        document.add_heading("Revenue", level=2)  # type: ignore[attr-defined]

        top = document.add_paragraph("North America", style="List Bullet")  # type: ignore[attr-defined]
        _set_numbering(top, 0)
        nested = document.add_paragraph("Enterprise accounts", style="List Bullet")  # type: ignore[attr-defined]
        _set_numbering(nested, 1)

        table = document.add_table(rows=3, cols=3)  # type: ignore[attr-defined]
        labels = [
            ("Item", "Qty", "Notes"),
            ("Apple", "3", "fresh"),
            ("Pear", "5", "ripe"),
        ]
        for row, values in zip(table.rows, labels, strict=True):
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = value

    return _docx_bytes(build)


def _docx_source(content: bytes, *, name: str = "fixture.docx") -> RawSource:
    return RawSource(name=name, content=content, media_type=_DOCX_MEDIA_TYPE)


# --------------------------------------------------------------------------- #
# parse_docx: a single layout-free page of Markdown (acceptance criterion 1)
# --------------------------------------------------------------------------- #


def test_docx_parses_to_single_layout_free_page() -> None:
    pages = parse_docx(_docx_source(_structured_docx()), _structured_docx())
    assert len(pages) == 1
    page = pages[0]
    assert isinstance(page, Page)
    assert page.number == 1
    assert page.layout is None  # flow format: classifier falls back to text-presence
    assert page.image is None
    assert page.text.strip()  # non-blank, so it routes as TEXT


# --------------------------------------------------------------------------- #
# parse_docx: styles/lists/tables → Markdown (acceptance criterion 2)
# --------------------------------------------------------------------------- #


def test_heading_styles_become_atx_headings() -> None:
    text = parse_docx(_docx_source(_structured_docx()), _structured_docx())[0].text
    assert "# Quarterly Report" in text
    assert "## Revenue" in text


def test_title_maps_to_h1() -> None:
    def build(document: object) -> None:
        document.add_heading("Doc Title", level=0)  # type: ignore[attr-defined]  # level 0 -> Title style

    content = _docx_bytes(build)
    text = parse_docx(_docx_source(content), content)[0].text
    assert text.startswith("# Doc Title")


def test_list_levels_are_preserved_with_indentation() -> None:
    text = parse_docx(_docx_source(_structured_docx()), _structured_docx())[0].text
    lines = text.splitlines()
    assert "- North America" in lines  # top level: no indent
    assert "  - Enterprise accounts" in lines  # ilvl=1: two-space indent


def test_numbered_list_uses_ordered_marker() -> None:
    def build(document: object) -> None:
        para = document.add_paragraph("First step", style="List Number")  # type: ignore[attr-defined]
        _set_numbering(para, 0)

    content = _docx_bytes(build)
    text = parse_docx(_docx_source(content), content)[0].text
    assert "1. First step" in text.splitlines()


def test_table_grid_becomes_pipe_table() -> None:
    text = parse_docx(_docx_source(_structured_docx()), _structured_docx())[0].text
    assert "| Item | Qty | Notes |" in text
    assert "| --- | --- | --- |" in text
    assert "| Apple | 3 | fresh |" in text


def test_pipe_in_cell_is_escaped() -> None:
    def build(document: object) -> None:
        table = document.add_table(rows=2, cols=1)  # type: ignore[attr-defined]
        table.rows[0].cells[0].text = "Header"
        table.rows[1].cells[0].text = "a | b"

    content = _docx_bytes(build)
    text = parse_docx(_docx_source(content), content)[0].text
    assert "| a \\| b |" in text  # the literal pipe is escaped, not a new column


def test_blocks_stay_in_document_order() -> None:
    # Heading, then list, then table — the interleaved body walk must preserve it.
    text = parse_docx(_docx_source(_structured_docx()), _structured_docx())[0].text
    assert text.index("## Revenue") < text.index("- North America")
    assert text.index("- North America") < text.index("| Item | Qty | Notes |")


# --------------------------------------------------------------------------- #
# Through the real backend: heading metadata + whole tables (criterion 3)
# --------------------------------------------------------------------------- #


def test_chunking_carries_heading_context() -> None:
    backend = HeuristicBackend()
    source = _docx_source(_structured_docx())
    parsed = backend.parse(source)
    assert len(parsed.pages) == 1  # DOCX routes to the docx path and stays one page

    chunks = backend.chunk(parsed.document, parsed.pages, size=512, overlap=64)
    assert chunks
    headings = {tuple(chunk.metadata.get("headings", [])) for chunk in chunks}
    # Some chunk sits under the H1 > H2 path the styles defined.
    assert any("Quarterly Report" in path and "Revenue" in path for path in headings)


def test_table_survives_chunking_whole() -> None:
    backend = HeuristicBackend()
    source = _docx_source(_structured_docx())
    parsed = backend.parse(source)
    chunks = backend.chunk(parsed.document, parsed.pages, size=512, overlap=64)

    table_chunks = [c for c in chunks if "| Item | Qty | Notes |" in c.text]
    assert len(table_chunks) == 1  # the whole grid lands in exactly one chunk
    table_text = table_chunks[0].text
    assert "| Apple | 3 | fresh |" in table_text
    assert "| Pear | 5 | ripe |" in table_text
