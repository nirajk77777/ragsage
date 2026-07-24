"""The DOCX parse path — honour Word's real structure, don't guess it.

Where the PDF path must *infer* structure from geometry, a ``.docx`` carries it
explicitly: heading styles, list levels, and table grids are all recorded in the
OOXML. So this path reads the source rather than second-guessing it — real
heading styles become Markdown ATX headings, list nesting becomes indented
Markdown list items, and table grids become GitHub pipe tables. The shared
two-pass chunker (:mod:`ragsage.parsing.chunking`) then splits on those headings,
carries the heading path onto each chunk, and keeps every table whole.

DOCX is a *flow* format: there is no page grid to measure, so the whole document
collapses onto a single :class:`~ragsage.models.Page` with ``layout=None`` — the
same contract the Docling adapter used for flow formats, which makes the page
classifier fall back to text-presence and route the page as ``TEXT``.

``python-docx`` is imported **lazily inside** :func:`parse_docx` (as ``_docx``,
to avoid any confusion with this module's own name ``ragsage.parsing.docx``) so
that importing :mod:`ragsage.parsing` never drags a parser library in and the
CPU-compatibility dependency guard stays green.

One subtlety drives the whole walk: iterating ``document.paragraphs`` and
``document.tables`` separately *loses their interleaving order*. To emit blocks in
document order we walk the body's XML children directly, dispatching ``<w:p>`` to
a paragraph and ``<w:tbl>`` to a table, so a heading, the list under it, and the
table after it all come out where the author put them.
"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING

from ragsage.models import Page, RawSource

if TYPE_CHECKING:  # pragma: no cover - typing only, never imported at runtime
    from docx.table import Table
    from docx.text.paragraph import Paragraph

# Flow format: no page grid, so the document is one page and the classifier's
# text-presence fallback routes it as TEXT.
_FLOW_FORMAT_PAGE = 1

# Markdown only splits headings up to level 6 (### #### the chunker cares about);
# a deeper Word heading is clamped rather than emitting a runaway run of hashes.
_MAX_HEADING_LEVEL = 6

# 2 spaces of indent per list nesting level, so Markdown reflects Word's ilvl.
_INDENT_PER_LEVEL = "  "

# A Word list paragraph style, optionally suffixed with a 1-based nesting level
# ("List Bullet 2" is the second level). ``List Paragraph`` is Word's generic
# list style; it counts as a list only when it also carries numbering (see below).
_LIST_STYLE = re.compile(r"^(List Bullet|List Number|List Paragraph)(?:\s+(\d+))?$", re.IGNORECASE)

# A real heading style: "Heading N", or "Title" (the document title → H1).
_HEADING_STYLE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def parse_docx(source: RawSource, content: bytes) -> list[Page]:
    """Parse DOCX ``content`` into a single Markdown :class:`~ragsage.models.Page`.

    ``source`` is accepted for parity with the other format paths (its name is the
    document label upstream); the bytes are the authority and are read here. The
    returned page is layout-free, so the classifier routes it as ``TEXT``.
    """
    import docx as _docx  # lazy: keeps ``import ragsage.parsing`` parser-free
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = _docx.Document(io.BytesIO(content))

    blocks: list[str] = []
    pending_list: list[str] = []

    def flush_list() -> None:
        # Consecutive list items form one Markdown list (single-newline joined) so
        # the levels read as one nested list, not a run of stray blocks.
        if pending_list:
            blocks.append("\n".join(pending_list))
            pending_list.clear()

    # Walk the body's children in document order (see module docstring); dispatch
    # by tag so paragraphs, lists, and tables stay interleaved as authored.
    p_tag, tbl_tag = qn("w:p"), qn("w:tbl")
    for child in document.element.body.iterchildren():
        if child.tag == p_tag:
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue  # blank paragraph — a spacer, nothing to serialise
            style_name = _style_name(paragraph)
            list_line = _render_list_item(paragraph, style_name, text)
            if list_line is not None:
                pending_list.append(list_line)
                continue
            flush_list()
            blocks.append(_render_block(style_name, text))
        elif child.tag == tbl_tag:
            flush_list()
            table = _render_table(Table(child, document))
            if table:
                blocks.append(table)
    flush_list()

    markdown = "\n\n".join(blocks)
    return [Page(number=_FLOW_FORMAT_PAGE, text=markdown, layout=None)]


def _style_name(paragraph: Paragraph) -> str:
    """The paragraph's style name, or ``""`` when it has no resolvable style."""
    style = paragraph.style
    return style.name or "" if style is not None else ""


def _render_block(style_name: str, text: str) -> str:
    """Render a non-list paragraph: a heading style → ATX heading, else plain text."""
    level = _heading_level(style_name)
    if level is not None:
        return f"{'#' * level} {text}"
    return text


def _heading_level(style_name: str) -> int | None:
    """The Markdown heading level for a Word style, or ``None`` if it isn't one.

    ``Title`` maps to H1; ``Heading N`` maps to level ``N`` (clamped to 6).
    """
    if style_name.strip().lower() == "title":
        return 1
    match = _HEADING_STYLE.match(style_name.strip())
    if match is None:
        return None
    return min(int(match.group(1)), _MAX_HEADING_LEVEL)


def _render_list_item(paragraph: Paragraph, style_name: str, text: str) -> str | None:
    """Render a list paragraph as an indented Markdown item, or ``None`` if it isn't one.

    Membership and nesting are read from the paragraph's numbering properties
    (``<w:numPr>``/``<w:ilvl>``) first — the authoritative signal — falling back to
    the list style name (and its ``List Bullet 2`` style-level suffix). A numbered
    list uses a ``1.`` marker, everything else a ``-`` bullet.
    """
    numbering_level = _numbering_level(paragraph)
    style_match = _LIST_STYLE.match(style_name.strip())
    if numbering_level is None and style_match is None:
        return None

    level = numbering_level
    if level is None:
        # No numPr: derive the level from a "List Bullet 2" style suffix (1-based).
        assert style_match is not None
        suffix = style_match.group(2)
        level = int(suffix) - 1 if suffix else 0

    ordered = style_match is not None and style_match.group(1).lower() == "list number"
    marker = "1." if ordered else "-"
    return f"{_INDENT_PER_LEVEL * level}{marker} {text}"


def _numbering_level(paragraph: Paragraph) -> int | None:
    """The 0-based ``<w:ilvl>`` of a paragraph's numbering, or ``None`` if unnumbered.

    A ``<w:numPr>`` with no explicit ``<w:ilvl>`` is level 0. Read straight off the
    OOXML with qualified names so it works regardless of python-docx's convenience
    accessors.
    """
    from docx.oxml.ns import qn

    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    val = ilvl.get(qn("w:val"))
    return int(val) if val is not None else 0


def _render_table(table: Table) -> str:
    """Serialise a Word table grid to a GitHub pipe table.

    The first row is the header, followed by a ``| --- |`` delimiter sized to it,
    then one line per data row. Cell text is flattened (a multi-paragraph cell's
    newlines become spaces) and any ``|`` is escaped so it can't break the grid;
    the chunker keeps the whole table in a single chunk.
    """
    rows = table.rows
    if not rows:
        return ""
    header = [_cell_text(cell) for cell in rows[0].cells]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(_cell_text(cell) for cell in row.cells) + " |")
    return "\n".join(lines)


def _cell_text(cell: object) -> str:
    """One table cell as a single Markdown-safe line (newlines flattened, pipes escaped)."""
    text = " ".join(getattr(cell, "text", "").split())
    return text.replace("|", "\\|")
